import streamlit as st
import re
import io
from google import genai
from google.genai import types
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

st.set_page_config(page_title="AI Discharge Summarizer", page_icon="🏥", layout="wide")

st.title("🏥 AI Discharge Summarizer + Clinical Safety Layer")
st.markdown("Генериране на официална българска епикриза с автоматичен одит за критични стойности.")

# 1. Страничен панел
with st.sidebar:
    st.header("⚙️ Настройки")
    api_key = st.text_input("Gemini API Key:", type="password")
    model_name = st.selectbox("Избор на модел:", [
        'gemini-3.5-flash',
        'gemini-2.5-flash',
        'gemini-3.1-pro-preview',
        'gemini-1.5-flash'
    ])
    st.info("Въведете своя Gemini API ключ, за да стартирате системата.")

# 2. Модули за анонимизация и одит
class MedicalAnonymizer:
    def __init__(self):
        self.egn_pattern = r'\b[0-9]{10}\b'
        self.phone_pattern = r'(\+359|0)\s*(8[7-9]|9[8-9]|2)\s*\d{3}\s*\d{3,4}'

    def anonymize(self, text: str):
        mapping = {}
        anonymized_text = text

        raw_phones = [m.group(0) for m in re.finditer(self.phone_pattern, anonymized_text)]
        for i, phone in enumerate(set(raw_phones), 1):
            token = f"[ТЕЛЕФОН_{i}]"
            mapping[token] = phone
            anonymized_text = anonymized_text.replace(phone, token)

        egns = re.findall(self.egn_pattern, anonymized_text)
        for i, egn in enumerate(set(egns), 1):
            token = f"[ЕГН_{i}]"
            mapping[token] = egn
            anonymized_text = anonymized_text.replace(egn, token)

        patient_match = re.search(r'(?:Пациентът|Пациент|Пациентка|Болен|Болна)\s*:?\s*([А-ЯБВГДЕЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЬЮЯ][а-ябвгдежзийклмнопрстуфхцчшщъьюя]+\s+[А-ЯБВГДЕЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЬЮЯ][а-ябвгдежзийклмнопрстуфхцчшщъьюя]+\s+[А-ЯБВГДЕЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЬЮЯ][а-ябвгдежзийклмнопрстуфхцчшщъьюя]+|[А-ЯБВГДЕЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЬЮЯ][а-ябвгдежзийклмнопрстуфхцчшщъьюя]+\s+[А-ЯБВГДЕЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЬЮЯ][а-ябвгдежзийклмнопрстуфхцчшщъьюя]+)', anonymized_text)

        if patient_match:
            full_patient_name = patient_match.group(1)
            token = "[ПАЦИЕНТ_ИМЕ]"
            mapping[token] = full_patient_name
            anonymized_text = anonymized_text.replace(full_patient_name, token)

        return anonymized_text, mapping

class SafetyChecker:
    @staticmethod
    def analyze_critical_labs(text: str):
        critical_alerts = []
        lab_pattern = r'([A-Za-zА-Яа-я0-9\s\+\-]+)\s*:?\s*([\d\.]+)\s*(\([^\)]+\)|[A-Za-z\/^0-9]+)?'
        
        if re.search(r'(калий|k\+)\s*:?\s*([6-9]\.[\d]+|[6-9])', text, re.IGNORECASE):
            critical_alerts.append("🚨 **КРИТИЧНА СТОЙНОСТ:** Хиперкалиемия ($K^+ > 6.0$ mmol/L)! Висок риск от аритмии. Проверете за АСЕ-инхибитори.")
        if re.search(r'(тропонин|troponin)\s*:?\s*([0-9]+\.[\d]+)', text, re.IGNORECASE):
            critical_alerts.append("🚨 **КРИТИЧНА СТОЙНОСТ:** Повишен Тропонин! Задължително потвърдете за остър миокарден инфаркт.")
        if re.search(r'(креатинин|creatinine)\s*:?\s*([1-9][0-9]{2})', text, re.IGNORECASE):
            critical_alerts.append("⚠️ **ОТКЛОНЕНИЕ:** Повишен Креатинин (Бъбречна недостатъчност/дисфункция).")
            
        return critical_alerts

def create_hospital_docx(summary_text: str) -> io.BytesIO:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    p_hospital = doc.add_paragraph()
    p_hospital.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_hospital.add_run("МНОГОПРОФИЛНА БОЛНИЦА ЗА АКТИВНО ЛЕЧЕНИЕ И КЛИНИКА\n")
    r1.bold = True
    r1.font.size = Pt(9)

    table_top = doc.add_table(rows=1, cols=2)
    table_top.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_top.cell(0, 0).paragraphs[0].add_run("Е П И К Р И З А").bold = True
    table_top.cell(0, 1).paragraphs[0].add_run("Клиника по Кардиология\nВх.№ 11897 / 2026г.")

    for line in summary_text.split('\n'):
        if line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(line.strip())
            r.font.name = "Times New Roman"
            r.font.size = Pt(10)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# 3. UI Компоненти
col1, col2 = st.columns([1, 1])

uploaded_file_bytes = None
uploaded_file_mime = None
extracted_text = ""

with col1:
    st.subheader("1. Входящи данни")
    uploaded_file = st.file_uploader("Качете PDF, изображение или TXT файл:", type=['pdf', 'png', 'jpg', 'jpeg', 'txt'])
    
    if uploaded_file is not None:
        file_ext = uploaded_file.name.split('.')[-1].lower()
        content = uploaded_file.read()
        
        if file_ext == 'txt':
            extracted_text = content.decode('utf-8', errors='ignore')
        elif file_ext == 'pdf':
            try:
                pdf_reader = PdfReader(io.BytesIO(content))
                for page in pdf_reader.pages:
                    t = page.extract_text()
                    if t:
                        extracted_text += t + "\n"
                if not extracted_text.strip():
                    extracted_text = f"[Сканиран PDF: {uploaded_file.name}]"
                    uploaded_file_bytes = content
                    uploaded_file_mime = "application/pdf"
            except Exception as e:
                st.error(f"Грешка: {e}")
        elif file_ext in ['png', 'jpg', 'jpeg']:
            extracted_text = f"[Снимка/Изображение: {uploaded_file.name}]"
            uploaded_file_bytes = content
            uploaded_file_mime = f"image/{'jpeg' if file_ext in ['jpg', 'jpeg'] else 'png'}"

    input_text = st.text_area("Текст на епикризата / изследванията:", value=extracted_text, height=280)

with col2:
    st.subheader("2. Обработка & Резултат")
    if st.button("🚀 Генерирай Епикриза", type="primary"):
        if not api_key:
            st.warning("⚠️ Въведете Gemini API Key в страничното меню!")
        elif not input_text.strip() and not uploaded_file_bytes:
            st.warning("⚠️ Въведете текст или качете файл!")
        else:
            with st.spinner("Генериране и деанонимизиране..."):
                try:
                    anonymizer = MedicalAnonymizer()
                    clean_text, data_map = anonymizer.anonymize(input_text)

                    system_instruction = """
                    Ти си висококвалифициран български лекар-ординатор.
                    Генерирай академична, безупречна ОФИЦИАЛНА ЕПИКРИЗА въз основа на предоставения текст.
                    
                    ЗАДЪЛЖИТЕЛНИ СЕКЦИИ С УДЕБЕЛЕНИ ЗАГЛАВИЯ:
                    1. ИДЕНТИФИКАЦИЯ НА ПАЦИЕНТА (Задължително запази токените [ПАЦИЕНТ_ИМЕ], [ЕГН_1] и т.н.)
                    2. ОКОНЧАТЕЛНА ДИАГНОЗА
                    3. АНАМНЕЗА И ОБЕКТИВНО СЪСТОЯНИЕ
                    4. ПАРАКЛИНИЧНИ ИЗСЛЕДВАНИЯ (Обърни внимание на патологичните стойности ⚠️)
                    5. КЛИНИЧЕН ХОД И ЕВОЛЮЦИЯ
                    6. ТЕРАПИЯ ПРИ ИЗПИСВАНЕ И ПРЕПОРЪКИ
                    """

                    client = genai.Client(api_key=api_key)
                    if uploaded_file_bytes:
                        part = types.Part.from_bytes(data=uploaded_file_bytes, mime_type=uploaded_file_mime)
                        contents = [part, "Извърши OCR и генерирай пълна епикриза:"]
                    else:
                        contents = [clean_text]

                    response = client.models.generate_content(
                        model=model_name,
                        contents=contents,
                        config={'system_instruction': system_instruction}
                    )
                    
                    ai_summary = response.text

                    # ДЕАНОНИМИЗАЦИЯ (Връщане на реалните имена)
                    final_summary = ai_summary
                    for token, real_val in data_map.items():
                        final_summary = final_summary.replace(token, real_val)

                    st.session_state['result'] = final_summary
                    st.session_state['alerts'] = SafetyChecker.analyze_critical_labs(input_text)
                except Exception as e:
                    st.error(f"Грешка: {e}")

    if 'result' in st.session_state:
        st.text_area("Генериран текст:", value=st.session_state['result'], height=280)
        
        alerts = st.session_state.get('alerts', [])
        if alerts:
            st.markdown("### 🛡️ Clinical Safety Audit")
            for a in alerts:
                st.warning(a)
        else:
            st.success("✅ Clinical Safety Check: Няма открити критични конфликти.")

        docx_bytes = create_hospital_docx(st.session_state['result'])
        st.download_button(
            label="💾 Свали ОФИЦИАЛНА БЛАНКА (.docx)",
            data=docx_bytes,
            file_name="Official_Hospital_Epikriza.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
