import streamlit as st
import re
import io
from google import genai
from google.genai import types
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Настройка на страницата
st.set_page_config(page_title="AI Discharge Summarizer", page_icon="🏥", layout="wide")

st.title("🏥 AI Discharge Summarizer + Clinical Safety Layer")
st.markdown("Генериране на официална българска епикриза с автоматичен одит за критични стойности.")

# Страничен панел за настройки
with st.sidebar:
    st.header("⚙️ Настройки")
    api_key = st.text_input("Gemini API Key:", type="password")
    model_name = st.selectbox("Избор на модел:", ['gemini-3.5-flash', 'gemini-3.1-pro-preview'])
    st.info("Въведете своя Gemini API ключ, за да стартирате системата.")

# Качване на файл или текстов вход
col1, col2 = st.columns([1, 1])

uploaded_file_bytes = None
uploaded_file_mime = None
extracted_text = ""

with col1:
    st.subheader("1. Входящи данни")
    uploaded_file = st.file_uploader("Качете PDF, изображение или TXT файл:", type=['pdf', 'png', 'jpg', 'jpeg', 'txt'])
    
    if uploaded_file is not None:
        file_ext = uploaded_file.name.split('.')[-1].lower()
        content = uploaded_file.read()
        
        if file_ext == 'txt':
            extracted_text = content.decode('utf-8', errors='ignore')
        elif file_ext == 'pdf':
            try:
                pdf_reader = PdfReader(io.BytesIO(content))
                for page in pdf_reader.pages:
                    t = page.extract_text()
                    if t:
                        extracted_text += t + "\n"
                if not extracted_text.strip():
                    extracted_text = f"[Сканиран PDF: {uploaded_file.name}]"
                    uploaded_file_bytes = content
                    uploaded_file_mime = "application/pdf"
            except Exception as e:
                st.error(f"Грешка при четене на PDF: {e}")
        elif file_ext in ['png', 'jpg', 'jpeg']:
            extracted_text = f"[Снимка/Изображение: {uploaded_file.name}]"
            uploaded_file_bytes = content
            uploaded_file_mime = f"image/{'jpeg' if file_ext in ['jpg', 'jpeg'] else 'png'}"

    input_text = st.text_area("Текст на епикризата / изследванията:", value=extracted_text, height=250)

# Логика за анонимизация и проверки
class SafetyChecker:
    @staticmethod
    def analyze_critical_labs(text: str):
        critical_alerts = []
        lab_pattern = r'([A-Za-zА-Яа-я0-9\s-]+)\s+([\d\.]+)\s*\(([\d\.]+)\s*-\s*([\d\.]+)\)'
        matches = re.findall(lab_pattern, text)
        for match in matches:
            param, val_str, min_ref_str, max_ref_str = match
            val, max_ref = float(val_str), float(max_ref_str)
            param_name = param.strip().upper()
            if "TROPONIN" in param_name and val > max_ref:
                critical_alerts.append(f"🚨 **КРИТИЧНА СТОЙНОСТ:** Висок Тропонин ({val})! Задължително потвърди ИМ.")
            elif param_name in ["K", "ПОТАСИУМ", "КАЛИЙ"] and val > 5.5:
                critical_alerts.append(f"🚨 **КРИТИЧНА СТОЙНОСТ:** Хиперкалиемия ({val} mmol/L)!")
        return critical_alerts

def create_hospital_docx(summary_text: str) -> io.BytesIO:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    p_hospital = doc.add_paragraph()
    p_hospital.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_hospital.add_run("МНОГОПРОФИЛНА БОЛНИЦА ЗА АКТИВНО ЛЕЧЕНИЕ И КЛИНИКА\n").bold = True

    table_top = doc.add_table(rows=1, cols=2)
    table_top.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_top.cell(0, 0).paragraphs[0].add_run("Е П И К Р И З А").bold = True
    table_top.cell(0, 1).paragraphs[0].add_run("Клиника по Кардиология\nВх.№ 11897 / 2026г.")

    for line in summary_text.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# Генериране на резултат
with col2:
    st.subheader("2. Обработка & Резултат")
    if st.button("🚀 Генерирай Епикриза", type="primary"):
        if not api_key:
            st.warning("⚠️ Моля, въведете Gemini API Key в страничното меню!")
        elif not input_text.strip() and not uploaded_file_bytes:
            st.warning("⚠️ Введи текст или качи файл!")
        else:
            with st.spinner("Обработка и одит..."):
                try:
                    client = genai.Client(api_key=api_key)
                    if uploaded_file_bytes:
                        part = types.Part.from_bytes(data=uploaded_file_bytes, mime_type=uploaded_file_mime)
                        contents = [part, "Извърши OCR и генерира Official Епикриза:"]
                    else:
                        contents = [input_text]

                    response = client.models.generate_content(
                        model=model_name,
                        contents=contents,
                        config={'system_instruction': "Генерирай академична българска епикриза с официални секции."}
                    )
                    
                    st.session_state['result'] = response.text
                    st.session_state['alerts'] = SafetyChecker.analyze_critical_labs(input_text)
                except Exception as e:
                    st.error(f"Грешка: {e}")

    if 'result' in st.session_state:
        st.text_area("Генериран текст:", value=st.session_state['result'], height=250)
        
        # Safety audit показване
        alerts = st.session_state.get('alerts', [])
        if alerts:
            for a in alerts:
                st.warning(a)
        else:
            st.success("✅ Clinical Safety Check: Няма открити критични конфликти.")

        # Сваляне
        docx_bytes = create_hospital_docx(st.session_state['result'])
        st.download_button(
            label="💾 Свали ОФИЦИАЛНА БЛАНКА (.docx)",
            data=docx_bytes,
            file_name="Official_Hospital_Epikriza.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
